import os,sys
import json
import PyPDF2
import docx
from pathlib import Path
import pandas as pd
from openai import OpenAI
from config.setting import CONFIG
import re
import requests
from datetime import datetime
import utils.word_to_images as word_to_images
import base64
import hashlib

def read_docx_by_paragraph(docx_path):
    """Read the DOCX content paragraph by paragraph."""
    doc = docx.Document(docx_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != '']
    return paragraphs

def split_paragraphs_into_pages(paragraphs, chars_per_page=1000):
    """Split paragraphs into pages by character count."""
    pages = []
    current_page = ""
    for para in paragraphs:
        if len(current_page) + len(para) + 1 > chars_per_page:
            pages.append(current_page)
            current_page = para
        else:
            current_page += "\n" + para
    if current_page:
        pages.append(current_page)
    return pages

def read_docx_by_page(docx_path, chars_per_page=1000):
    """Read DOCX and split content into pages by paragraphs."""
    paragraphs = read_docx_by_paragraph(docx_path)
    pages = split_paragraphs_into_pages(paragraphs, chars_per_page)
    return pages


def read_file_content(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return read_pdf(file_path)
    elif ext == '.docx':
        return read_docx_by_page(file_path,2000)
    elif ext == '.doc':
        return read_docx_by_page(file_path,2000)
    elif ext == '.csv':
        return read_csv(file_path)
    elif ext == '.xlsx':
        return read_xlsx(file_path)
    elif ext == '.txt':
        return read_txt(file_path)
    else:
        return None
    
def read_txt(file_path):
    """读取文件内容"""
    # with open(file_path, 'r', encoding='utf-8') as file:
    #     return file.readlines()
    pdf_path = word_to_images.convert_to_pdf(os.path.abspath(os.path.dirname(file_path)), os.path.basename(file_path))
    return read_pdf(pdf_path)

def read_pdf(file_path):
    content = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        page_numbers = len(reader.pages)
        for page_num in range(page_numbers):
            page = reader.pages[page_num]
            content.append(page.extract_text())
    return content

def read_docx(file_path):
    doc = docx.Document(file_path)
    return [para.text for para in doc.paragraphs]

def read_csv(file_path):
    df = pd.read_csv(file_path)
    json_str = df.to_json(orient='records')
    return json.loads(json_str)

def read_xlsx(file_path):
    df = pd.read_excel(file_path)
    json_str = df.to_json(orient='records')
    return json.loads(json_str)


def paginate_text(text, page_length=1000):
    """按页分页文本，每页指定字数，如果段落被截断则补充完整"""
    words = text.split()
    pages = []
    current_page = []
    current_length = 0
    
    for word in words:
        current_length += len(word) + 1  # 加1是为了包括空格或换行符
        current_page.append(word)
        
        if current_length >= page_length:
            # 检查是否截断段落
            if not word.endswith(('.', '!', '?', '"', '”')):  # 简单检查句子结尾标点
                while words and not words[0].endswith(('.', '!', '?', '"', '”')):
                    current_page.append(words.pop(0))
            
            pages.append(' '.join(current_page))
            current_page = []
            current_length = 0
    
    if current_page:
        pages.append(' '.join(current_page)) 
    return pages

def get_llm_response(prompt):
    try:
        client = OpenAI(
            api_key=CONFIG["openai_config"]['api_key'],
            base_url=CONFIG["openai_config"]['base_url'],
        )
        completion = client.chat.completions.create(
            model=CONFIG["openai_config"]['model'],
            messages=[
                {'role': 'user', 'content': prompt}
            ],
            extra_body={
                "top_k": 1,
            }
        )
        model_response = completion.dict()
        # model_response = completion.model_dump_json().encode("utf-8").decode("unicode_escape")
        # print(model_response)
        return model_response["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"get_response Error: {e}")


# 定义一个用于自然排序的函数
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split('([0-9]+)', s)]


class ServiceLogger:
    def __init__(self):
        self.url = "%s%s"%(CONFIG["app_log"]['host'],CONFIG["app_log"]['url'])
        self.log_type = "large"

    
    def app_logger(self, id,name,msg):
        try:
            if name == "大模型模板配置":
                self.log_type = "large"
            elif name == "数据前处理":
                self.log_type = "model-large-data"
            elif name == "算法基座":
                self.log_type == "model-large-base"
            data = {
                "id":id,
                "type":self.log_type,
                "message":msg
            }
            headers = {'Content-Type': 'application/json'}
            respose = requests.post(url=self.url,headers=headers, json=data,timeout=5)
            if respose.status_code != 200:
                print(respose.text)
            else:
                print("success")
        except Exception as e:
            print(e)
    
    @classmethod
    def procsss(cls,level,id,name,message):
        if CONFIG["app_log"]["enable"] == True:
            message_formatstr = f'{datetime.now()} - {level} - {message}'
            cls().app_logger(id,name,message_formatstr)
    
    @classmethod
    def info(cls,id,name,message):
        cls.procsss("__INFO__",id,name,message)

    @classmethod
    def error(cls,id,name,message):
        cls.procsss("__ERROR__",id,name,message)
    
    @classmethod
    def warning(cls,id,name,message):
        cls.procsss("__WARNING__",id,name,message)
    
    @classmethod
    def debug(cls, id,name,message):
        cls.procsss("__DEBUG__",id,name,message)


def workflow_llm_response(config, prompt):
    try:
        client = OpenAI(
            api_key=config["api_info"]['api_key'],
            base_url=config["api_info"]['api_url'],
        )
        if "model" in config["api_info"]:
            model = config["api_info"]['model']
        else:
            model = config['model']
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {'role': 'user', 'content': prompt}
            ],
            extra_body={
                "top_k": 1,
            }
        )
        model_response = completion.dict()
        return model_response["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"get_response Error: {e}")
   

def test_llm_response(model,config, prompt):
    try:
        client = OpenAI(
            api_key=config['api_key'],
            base_url=config['api_url'],
        )
        completion = client.chat.completions.create(
            # model=CONFIG["openai_config"]['model'],
            model = config['model'],
            messages=[
                {'role': 'user', 'content': prompt}
            ],
            extra_body={
                "top_k": 1,
            }
        )
        model_response = completion.dict()
        return model_response["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"get_response Error: {e}")
    

def parse_llm_info(entity_key,text):
    # text = "{提示词｜抽取结果｜结果所在的短句}"
    # 正则表达式模式，匹配 { 和 } 中间的三部分内容
    pattern = r'\{([^｜]+)｜([^｜]+)｜([^}]+)\}'
    # 使用 re.match() 函数匹配模式
    match = re.match(pattern, text)
    result = None
    if match:
        part1 = match.group(1)  # 提取第一部分内容
        part2 = match.group(2)  # 提取第二部分内容
        part3 = match.group(3)  # 提取第三部分内容
        if part1 == '提示词':
            return None
        if part2 == 'None':
            part2 = ""
        if part3 == 'None':
            part3 = ""
        result = {"entity_key":entity_key,"prompt":part1,"extract_content":part2,"origin_text":part3}
        return result
    else:
        print("匹配失败")
        return result

def make_hashable(obj):
    """
    将对象转换为可哈希的形式。
    支持嵌套列表、字典等。
    """
    print("obj:",obj)
    if isinstance(obj, (tuple, list)):
        # 将列表和元组转换为元组
        return tuple(make_hashable(item) for item in obj)
    elif isinstance(obj, dict):
        # 将字典转换为 frozenset
        return frozenset((key, make_hashable(value)) for key, value in obj.items())
    return obj  # 其他类型（如字符串、数字等）保持不变

def remove_duplicate_dicts(original_list):
    """
    移除列表中的重复字典。
    """
    seen = set()
    unique_list = []
    for d in original_list:
        if not isinstance(d, dict):  # 确保当前元素是字典
            continue
        # 将字典的值转换为可哈希形式
        hashable_d = make_hashable(d)
        if hashable_d not in seen:
            seen.add(hashable_d)
            unique_list.append(d)
    return unique_list
    
def filter_operators(input_string):
    try:
        # 定义运算符列表
        operators = [
            r'\+', r'\-', r'\*', r'\/', r'%', r'=', r'<', r'>', r'&', r'\|', r'\^', r'~', r'!',
            r'==', r'!=', r'>=', r'<=', r'<<', r'>>', r'\/\/=', r'\*\*=', r'and', r'or', r'not',
            r'is', r'is not', r'in', r'not in'
        ]
        # 将运算符列表拼接成一个正则表达式模式
        pattern = r'(' + r'|'.join(operators) + r')'
        # 使用 re.sub() 函数将匹配的运算符替换为空字符串
        filtered_string = re.sub(pattern, '', input_string)
        return filtered_string
    except Exception as e:
        print(e)
        return input_string
    

def filter_list(input_list):
    seen_empty = False
    filtered_list = []
    
    for item in input_list:
        if item['extract_content'] == '':
            if not seen_empty:
                filtered_list.append(item)
                seen_empty = True
        else:
            filtered_list.append(item)
    
    
    if len (filtered_list) == 1:
        return filtered_list 
    return [ item for item in filtered_list if item["extract_content"] != ""]

def clean_and_replace(text):
    # 去除字符串开头和结尾的|符号
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]

    # 替换字符串中连续的||为|
    while "||" in text:
        text = text.replace("||", "|")

    # 将字符串中的|替换成"或者"
    text = text.replace("|", "或者")

    return text

def find_project_root(current_path: Path, target_files: set) -> Path:
    """
    从当前路径开始，向上递归查找包含目标文件的目录，作为项目的根目录。

    Args:
        current_path (Path): 当前文件的路径。
        target_files (set): 目标文件或目录的集合。

    Returns:
        Path: 项目的根目录路径。
    """
    for parent in current_path.parents:
        if any((parent / target_file).exists() for target_file in target_files):
            return parent
    return current_path

def get_project_root() -> Path:
    """
    获取项目的根目录。

    Returns:
        Path: 项目的根目录路径。
    """
    current_file_path = Path(__file__).resolve()
    # 假设app.py在项目根目录
    project_root = find_project_root(current_file_path, {"app.py"})
    return project_root

def get_default_llm_config(user):
        default_llm_config_list = CONFIG["default_llm_list"]
        default_llm_config_list[0]["user"] = user
        default_llm_config_list[1]["user"] = user
        return default_llm_config_list

def get_default_ocr_config(user):     
        default_ocr_config_list = CONFIG["default_OCR_process_list"]
        default_ocr_config_list[0]["user"] = user
        default_ocr_config_list[1]["user"] = user
        return default_ocr_config_list


def get_pre_process_default_format():
    main_code = """
    def data_pre_process(input_data):
        '''
        对输入数据进行预处理 
        '''
        output_result = input_data

        #此处填写前处理逻辑 ...

        return output_result
    """    
    return main_code

def get_post_process_default_format():
    main_code = """
    def data_post_process(output_data):
        '''
        对输出数据进行预处理 
        '''
        output_result = output_data

        #此处填写后处理逻辑 ...

        return output_result
    """    
    return main_code

def file_to_base64(file_path):
    """
    将文件转换为 Base64 字符串
    :param file_path: 文件路径
    :return: Base64 编码字符串
    """
    try:
        with open(file_path, 'rb') as file:
            # 读取文件内容并进行 Base64 编码
            encoded_string = base64.b64encode(file.read()).decode('utf-8')
        return encoded_string
    except Exception as e:
        return f"Error: {e}"

def save_base64_to_file(base64_string, output_file):
    with open(output_file, 'w') as file:
        file.write(base64_string)


def get_md5(file_path, chunk_size=8192):
    """ 计算文件的 MD5 哈希值 """
    md5 = hashlib.md5()
    
    try:
        with open(file_path, "rb") as f:
            while chunk := f.read(chunk_size):  # 逐块读取文件
                md5.update(chunk)
    except FileNotFoundError:
        print(f"文件 {file_path} 不存在！")
        return None
    
    return md5.hexdigest()